from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "球赛购票系统需求分析与详细设计书.docx"


def set_run_font(run, name="宋体", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_format(paragraph, before=0, after=6, line=1.1, alignment=None):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    if alignment is not None:
        paragraph.alignment = alignment


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths[i] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[i]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, before=0, after=8, line=1.1)
    r = p.add_run(title)
    set_run_font(r, size=22, bold=True, color="0B2545")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, before=0, after=18, line=1.1)
    r = p.add_run(subtitle)
    set_run_font(r, size=12, color="555555")


def add_heading(doc, text, level):
    doc.add_heading(text, level=level)


def add_para(doc, text):
    p = doc.add_paragraph()
    set_paragraph_format(p)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph_format(p, after=4, line=1.1)
        r = p.add_run(item)
        set_run_font(r)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_paragraph_format(p, after=4, line=1.1)
        r = p.add_run(item)
        set_run_font(r)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "F2F4F7")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                set_run_font(run, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for p in cells[i].paragraphs:
                set_paragraph_format(p, after=0, line=1.05)
                for run in p.runs:
                    set_run_font(run, size=10.5)
    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def add_footer(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("足球联赛购票系统需求分析与详细设计书")
    set_run_font(r, size=9, color="666666")


def build():
    doc = Document()
    style_document(doc)
    add_footer(doc)

    add_title(doc, "足球联赛购票系统", "需求分析与系统详细设计书")
    add_para(
        doc,
        "本文档面向课程设计阶段二成果，围绕足球联赛购票系统的业务目标、角色、功能需求、非功能需求、"
        "系统架构、功能模块和数据库设计进行说明，为后续编码实现、测试和最终总报告撰写提供依据。"
        "系统默认采用 Spring Boot + Vue + MySQL + MyBatis 的技术路线。"
    )

    add_heading(doc, "1 系统需求", 1)
    add_heading(doc, "1.1 功能需求", 2)
    add_heading(doc, "1.1.1 系统角色及其功能", 3)
    add_para(
        doc,
        "足球联赛购票系统主要服务于有观赛需求的普通用户，同时为联赛统一票务平台提供赛季赛程、主客场比赛、场馆、票务和检票管理能力。"
        "系统参考英超等职业足球联赛的组织方式：多个俱乐部在一个赛季内进行主客场双循环比赛，每轮包含多场对阵。"
        "本系统面向全联赛所有俱乐部比赛售票，而不是单个俱乐部的主场票务系统。根据系统使用者的职责差异，可以将系统角色划分为用户、俱乐部、检票员和管理员。"
    )
    add_table(
        doc,
        ["角色", "角色说明", "主要功能"],
        [
            [
                "用户",
                "浏览足球联赛完整赛程并购买任意可售比赛门票的普通用户。",
                "注册登录、完善个人信息、浏览赛季赛程、按俱乐部或轮次筛选比赛、查看票区票价和剩余连坐能力、选择票区和购票张数、提交订单、模拟支付、查看电子票、申请退票、查看订单记录。",
            ],
            [
                "俱乐部",
                "联赛中的参赛俱乐部，拥有本俱乐部业务后台账号，负责维护本俱乐部和主场相关资料。",
                "维护俱乐部名称、队标、主场城市、主场地址、球员名单、教练信息、赛季战绩、主场座位布局和主场票区方案，查看本俱乐部主场比赛售票情况。",
            ],
            [
                "检票员",
                "负责比赛入场时核验电子票状态的工作人员。",
                "扫描或输入电子票码、核验订单和场次信息、标记已入场、识别无效票和重复入场。",
            ],
            [
                "管理员",
                "负责维护全联赛基础数据、发布赛程、管理统一票务和维护系统运行的管理人员。",
                "维护用户和角色权限、维护赛季和轮次、统一发布主客场比赛、导入或维护比赛成绩、维护俱乐部和主场信息、设置票区票价、管理座位库存、处理退票申请、查看全联赛销售统计、查看日志和维护系统参数。",
            ],
        ],
        [1300, 2600, 5460],
    )

    add_heading(doc, "1.1.2 需求规定", 3)
    add_para(
        doc,
        "系统功能围绕“浏览赛事、选择票区、填写购票张数、系统自动分配连坐座位、生成订单、完成支付、生成电子票、现场检票、后台管理”这一完整购票链路展开。"
        "各功能点既要支持用户端操作，也要保证后台对赛事和票务资源的可维护性。"
    )
    add_table(
        doc,
        ["功能编号", "功能名称", "功能说明"],
        [
            ["F01", "用户注册与登录", "普通用户和俱乐部负责人可自主注册；赛事管理员和系统管理员由已有系统管理员内部先行注册。登录时需选择身份，管理员登录需额外校验工号。"],
            ["F02", "赛事浏览与查询", "用户可按赛季、轮次、主队、客队、比赛时间、场馆、票价区间等条件查询比赛，并查看比赛详情。"],
            ["F03", "票区与连坐能力展示", "系统按比赛票区展示票价、剩余票数和当前最大可接纳连坐数；未最终确定的行不显示具体座位，用户只选择票区和购票张数。"],
            ["F04", "订单创建与座位规划", "用户确认比赛、票区和张数后创建订单；系统限制单笔最多购买 4 张，并通过座位分配算法尽量规划连续座位，然后锁定对应库存。"],
            ["F05", "模拟支付", "用户在规定时间内完成模拟支付；支付成功后订单状态变为已支付，系统生成电子票。"],
            ["F06", "电子票管理", "系统为每张票生成唯一电子票码，用户可查看电子票信息、入场时间、座位和检票状态。"],
            ["F07", "退票申请", "用户在比赛开始前按规则提交退票申请；管理员审核后释放座位并更新订单状态。"],
            ["F08", "检票核验", "检票员通过电子票码核验票据，系统校验场次、票据状态和入场次数，防止重复入场。"],
            ["F09", "俱乐部资料与主场管理", "俱乐部维护本俱乐部名称、队标、主场城市、主场地址、球员、教练、赛季胜场、主场座位布局和票区方案。"],
            ["F10", "赛事与票务后台管理", "管理员统一维护赛季、轮次、俱乐部、主客场比赛、比赛成绩、主场座位布局、场馆、票区票价、座位库存和上下架状态。"],
            ["F11", "统计分析", "系统统计全联赛、各轮次、各俱乐部主场比赛的售票数量、销售额、退票数量、上座率和热门赛事，为运营管理提供依据。"],
        ],
        [1200, 1800, 6360],
    )

    add_heading(doc, "1.1.3 业务流程概述", 3)
    add_numbered(
        doc,
        [
            "俱乐部维护本俱乐部名称、队标、主场城市、主场地址、球员名单、教练信息、赛季战绩和主场座位布局，作为后续主场比赛展示和票务配置的基础数据。",
            "管理员创建赛季和轮次，基于俱乐部和主场数据维护全联赛主客场赛程，并将比赛发布为可售状态；比赛结束后由管理员导入或维护比赛成绩。",
            "俱乐部可维护本俱乐部主场票区方案和建议票价，管理员统一确认正式售票配置。",
            "购票用户登录系统后浏览联赛赛程，按轮次或俱乐部进入比赛详情页，查看主客队、场馆、票区票价、余票数量和该票区当前最大可接纳连坐数。",
            "在座位行尚未最终确定前，用户端不展示具体座位号，只展示票区、余票和最大连坐能力，以便系统继续按订单组合优化同一排内的连坐规划。",
            "用户选择票区并填写购票张数后提交订单，系统校验张数是否超过单笔上限 4 张，并调用座位分配算法优先规划连续座位。",
            "若存在满足张数的连续座位，系统直接分配该组座位；若没有完整连坐但仍有余票，系统提示用户当前最大可连坐数，并建议用户分开下单。",
            "订单生成后，系统在待支付期间临时锁定规划出的座位或行内候选位置，避免其他用户重复购买。",
            "用户完成模拟支付后，系统确认订单为已支付状态，扣减库存并为每张票生成唯一电子票码。",
            "比赛入场时，检票员扫描或输入电子票码，系统核验票据是否属于当前场次且未被使用。",
            "若票据有效，系统记录检票时间并将电子票状态更新为已入场；若票据无效或重复使用，系统提示拒绝入场。",
        ],
    )

    add_heading(doc, "1.1.4 业务实体分析", 3)
    add_para(
        doc,
        "根据足球联赛购票业务，系统涉及的核心实体包括用户、角色、赛季、轮次、比赛、比赛成绩、俱乐部、球员、球员赛季数据、教练、赛季战绩、场馆、票区、座位、订单、电子票、支付记录、退票申请和检票记录。"
        "这些实体共同支撑从赛事发布到购票入场的完整数据链路。"
    )
    add_table(
        doc,
        ["实体", "关键属性", "说明"],
        [
            ["用户", "用户编号、姓名、手机号、密码、角色、所属俱乐部、状态、创建时间", "记录用户、俱乐部、检票员和管理员等账号基础信息。"],
            ["角色", "角色编号、角色名称、权限集合", "区分用户、俱乐部、检票员和管理员。"],
            ["赛季", "赛季编号、赛季名称、开始日期、结束日期、状态", "记录一个完整联赛赛季，是双循环赛程的组织单位。"],
            ["轮次", "轮次编号、赛季编号、轮次名称、开始日期、结束日期", "记录联赛第几轮比赛，便于赛程展示和查询。"],
            ["俱乐部", "俱乐部编号、俱乐部名称、队标、主场城市、主场地址、简介", "记录参赛俱乐部展示信息；俱乐部注册时需要录入主场相关信息。"],
            ["球员", "球员编号、俱乐部编号、姓名、号码、位置、状态", "记录俱乐部球员名单，作为后续射手榜、转会等扩展功能的基础对象。"],
            ["球员赛季数据", "赛季编号、球员编号、出场数、进球数、助攻数、所属俱乐部", "记录球员在某个赛季的统计数据，可用于扩展射手榜等功能。"],
            ["教练", "教练编号、俱乐部编号、姓名、职务、简介", "记录俱乐部教练信息，用于俱乐部资料和比赛详情页展示。"],
            ["赛季战绩", "赛季编号、俱乐部编号、胜场、平场、负场、进球、失球、积分", "记录俱乐部在某个赛季的联赛表现，用于积分榜和比赛详情页展示。"],
            ["场馆", "场馆编号、场馆名称、城市、地址、容量、排数、座位布局说明", "记录俱乐部主场或比赛举办地点及其座位布局基础信息。"],
            ["比赛", "比赛编号、赛季、轮次、主队、客队、比赛时间、场馆、比分、状态", "描述一场具体足球比赛，是购票业务和比赛结果展示的核心对象。"],
            ["票区", "票区编号、赛事编号、区域名称、票价、总票数、剩余票数、最大连坐数", "按区域管理不同价格、库存和当前可连续分配能力。"],
            ["座位", "座位编号、场馆编号、票区编号、排号、座号、状态", "记录可售、锁定、已售、停用等座位状态，供系统自动规划连坐座位。"],
            ["订单", "订单编号、用户编号、总金额、订单状态、创建时间、支付截止时间", "记录用户一次购票行为及其状态流转。"],
            ["电子票", "票码、订单编号、比赛编号、座位编号、票状态、入场时间", "作为用户入场凭证，需保证唯一性和不可重复使用。"],
            ["支付记录", "支付编号、订单编号、支付金额、支付方式、支付状态、支付时间", "记录订单支付过程，课程设计中可实现模拟支付。"],
            ["退票申请", "申请编号、订单编号、申请原因、审核状态、审核时间", "支持用户退票和后台审核。"],
            ["检票记录", "记录编号、票码、检票员编号、检票结果、检票时间", "记录入场核验过程，便于追溯异常票据。"],
        ],
        [1500, 3600, 4260],
    )

    add_heading(doc, "1.2 非功能性需求", 2)
    add_table(
        doc,
        ["类别", "需求说明"],
        [
            ["性能需求", "常规查询页面响应时间应控制在 2 秒以内；购票高峰期订单创建、库存扣减和连坐座位规划应避免超卖和长时间等待。"],
            ["安全需求", "用户密码需加密存储；后台接口应按角色鉴权；电子票码应具有唯一性，防止伪造和重复使用。"],
            ["可靠性需求", "支付成功、库存扣减和电子票生成应保持事务一致；异常中断时订单状态应可恢复。"],
            ["易用性需求", "赛事列表、票区选择、购票张数填写、订单确认和电子票查看流程应清晰，用户不需要手动寻找座位即可完成购票。"],
            ["可维护性需求", "采用 MVC 分层结构，将控制层、业务逻辑层、数据访问层和前端展示层分离，便于扩展和测试。"],
            ["兼容性需求", "用户端页面适配常见桌面浏览器和移动端浏览器；后台管理页面优先保证桌面端操作效率。"],
        ],
        [1800, 7560],
    )

    add_heading(doc, "2 系统设计", 1)
    add_heading(doc, "2.1 系统总体架构设计", 2)
    add_para(
        doc,
        "系统按照使用场景划分为用户端、俱乐部端、检票端、管理员后台和数据与算法层。"
        "前四个模块面向不同角色提供业务功能，数据与算法层为各端提供统一的数据存储、票务状态管理和连坐规划能力。"
    )
    add_table(
        doc,
        ["模块", "主要功能", "面向角色"],
        [
            ["用户端", "完善个人信息、浏览赛程、查看俱乐部资料、查看比赛详情、选择票区购票、查看订单和电子票、提交退票申请。", "用户"],
            ["俱乐部端", "维护俱乐部资料、球员、教练、主场信息、主场座位布局、票区方案，查看本俱乐部主场销售情况。", "俱乐部"],
            ["检票端", "输入或扫描电子票码，核验票据状态，记录入场结果，识别无效票和重复入场。", "检票员"],
            ["管理员后台", "维护赛季、轮次、赛程、比赛成绩、票务配置、退票处理、用户权限和统计分析。", "管理员"],
            ["数据与算法层", "保存俱乐部、赛程、票务、订单、电子票和检票数据，提供连坐规划和最大连坐数计算。", "系统内部"],
        ],
        [1600, 5260, 2500],
    )
    add_para(
        doc,
        "技术实现上，系统采用前后端分离的 MVC 分层思想。前端 Vue 页面负责用户交互和数据展示；后端 Spring Boot 提供 REST API；"
        "Controller 层接收请求并进行参数校验；Service 层封装购票、支付、退票、检票、比赛成绩维护和连坐规划等核心业务逻辑；"
        "Mapper/DAO 层通过 MyBatis 访问 MySQL 数据库。"
    )
    add_table(
        doc,
        ["技术层次", "主要职责", "建议技术"],
        [
            ["表示层", "用户端、俱乐部端、检票端和管理员后台页面。", "Vue、Element Plus、Axios"],
            ["控制层", "接收前端请求，完成参数校验、登录鉴权、接口响应封装。", "Spring MVC Controller"],
            ["业务逻辑层", "处理连坐座位规划、库存锁定、订单支付、电子票生成、退票处理、检票核验和比赛成绩维护等核心规则。", "Spring Boot Service"],
            ["数据访问层", "完成用户、俱乐部、赛事、票务、订单、检票记录等数据的增删改查。", "MyBatis Mapper"],
            ["数据层", "存储系统核心业务数据，并通过事务保证订单、库存和电子票状态一致。", "MySQL"],
        ],
        [1600, 5260, 2500],
    )

    add_heading(doc, "2.2 功能模块设计", 2)
    add_heading(doc, "2.2.1 功能模块划分", 3)
    add_table(
        doc,
        ["模块", "子功能", "说明"],
        [
            ["账号权限模块", "注册、登录、个人信息、角色权限、所属俱乐部绑定", "统一管理用户身份和后台权限。"],
            ["赛事浏览模块", "赛事列表、条件筛选、赛事详情、俱乐部资料、场馆信息", "为购票用户提供赛事发现入口。"],
            ["购票订单模块", "票区选择、张数限制、连坐座位规划、创建订单、库存锁定、订单查询、取消订单", "完成购票核心流程。"],
            ["支付与电子票模块", "模拟支付、电子票生成、票码展示、票据状态管理", "连接订单和入场凭证。"],
            ["退票管理模块", "退票申请、退票审核、库存释放、订单状态更新", "支持比赛开始前的退票处理。"],
            ["检票入场模块", "票码核验、入场记录、重复检票识别", "保障现场入场秩序。"],
            ["俱乐部管理模块", "俱乐部名称、队标、主场城市、主场地址、球员、教练、赛季战绩、主场座位布局、票区方案维护", "支持各俱乐部维护自己的主场基础资料。"],
            ["后台运营模块", "赛季、轮次、俱乐部、主场座位布局、全联赛主客场比赛、比赛成绩、场馆、票区、票价、库存维护", "支撑统一联赛赛程发布、成绩导入和票务资源管理。"],
            ["统计分析模块", "销售额、售票量、退票量、上座率、赛季战绩统计", "为管理员提供运营数据和联赛展示数据。"],
        ],
        [1800, 3000, 4560],
    )

    add_heading(doc, "2.2.2 购票功能点设计", 3)
    add_para(
        doc,
        "购票功能是系统的核心功能点，涉及前端页面、后端接口、业务服务和数据库表的协同。该功能不要求用户手动选择具体座位，"
        "而是由用户选择票区和购票张数，系统根据该票区的座位分布与实时状态自动规划座位。不同俱乐部主场的具体排数和每排宽度可能不同，"
        "因此座位规划应基于场馆录入的实际座位布局数据执行，而不能假设每个票区拥有固定排数或固定每排座位数。该功能需要重点处理连坐分配、座位状态、订单状态和支付状态的一致性。"
    )
    add_table(
        doc,
        ["资源类型", "资源名称", "说明"],
        [
            ["前端页面", "MatchList.vue、MatchDetail.vue、ZoneSelect.vue、OrderConfirm.vue", "展示赛事、票区、当前最大连坐数、订单确认和支付入口。"],
            ["控制层", "MatchController、TicketController、OrderController", "提供赛事查询、票区查询、连坐能力查询、订单创建和支付确认接口。"],
            ["业务层", "MatchService、TicketService、SeatAllocateService、OrderService、PaymentService", "处理赛事读取、库存校验、连坐座位规划、订单创建、支付确认和电子票生成。"],
            ["数据访问层", "MatchMapper、SeatMapper、OrderMapper、TicketMapper", "访问赛事、座位、订单和电子票相关数据表。"],
            ["数据库表", "match_info、seat_info、ticket_zone、ticket_order、e_ticket、payment_record", "保存购票流程相关核心数据。"],
        ],
        [1600, 3300, 4460],
    )
    add_para(
        doc,
        "购票业务逻辑过程为：用户选择赛事后进入票区选择页面；系统查询该比赛下各票区的票价、余票和当前最大可接纳连坐数；"
        "用户选择票区并输入购票张数，张数不得超过系统设置的单笔上限 4 张；后端在事务中校验票区余票和座位状态，并调用连坐座位分配算法。"
        "系统以“排”为基本规划单位，连坐必须发生在同一排内。每个俱乐部主场的排数、每排座位数和票区划分均来自场馆座位布局数据。"
        "若某一排尚未排满，且后续订单仍可能改变该排内不同订单的位置组合，系统暂不向用户公布具体座位号；"
        "当某一排已经排满，或该排剩余空间已经无法形成更优组合时，该排座位规划被标记为最终确定，系统再向相关用户展示具体座位。"
        "算法按排号优先级扫描该票区各排的剩余连续空间，优先寻找能够容纳本次购票张数的行内连续区间；若存在多个可选行，则优先选择位置更靠前、剩余空间利用率更高、座位更靠近看台中线的行。"
        "越早下单的用户，在满足连坐和行级规划的前提下，系统应尽量分配更靠前、更居中的座位。"
        "以一个简化示例说明：若某票区每排 6 个座位、共两排，购票张数顺序为 4、1、2、1、3，系统可将第一排规划为 1、4、1，第二排规划为 2、3，此时剩余最大连续空间为 1，并提示用户最多可购买 1 张；"
        "若购票张数顺序为 4、2、1、1、3，则可将第一排规划为 4、2，第二排规划为 1、1、3。"
        "若找到合适座位，系统生成待支付订单并将座位状态改为锁定；用户支付成功后，系统将订单状态改为已支付，"
        "座位状态改为已售，并为每个座位生成电子票。若用户超时未支付，系统取消订单并释放座位。"
    )

    add_heading(doc, "2.2.3 检票功能点设计", 3)
    add_para(
        doc,
        "检票功能用于比赛现场入场核验。检票员输入或扫描电子票码后，系统查询电子票、订单、赛事和座位信息，判断票据是否有效。"
        "若票据属于当前场次、订单已支付且票据未使用，则允许入场并记录检票结果；否则提示票据无效、已退票、未支付或已入场。"
    )
    add_table(
        doc,
        ["校验项", "通过条件", "失败处理"],
        [
            ["票码真实性", "电子票码在系统中存在且未被删除。", "提示票码不存在或无效。"],
            ["订单状态", "关联订单状态为已支付。", "提示订单未支付、已取消或已退票。"],
            ["赛事匹配", "电子票对应赛事为当前检票场次。", "提示非当前场次票据。"],
            ["入场状态", "电子票状态为未使用。", "提示该票已入场，禁止重复检票。"],
        ],
        [1800, 3780, 3780],
    )

    add_heading(doc, "2.3 数据库设计", 2)
    add_heading(doc, "2.3.1 逻辑结构设计", 3)
    add_para(
        doc,
        "数据库逻辑结构围绕赛事、票务、订单和检票四条主线建立。用户通过订单购买赛事座位；订单支付成功后生成电子票；"
        "电子票在比赛入场时产生检票记录。主要实体关系如下："
    )
    add_bullets(
        doc,
        [
            "一个用户可以创建多个订单，一个订单只属于一个用户。",
            "一个赛季包含多个轮次，一个轮次包含多场比赛。",
            "一个俱乐部在一个赛季对应一条赛季战绩记录，用于展示胜场、平场、负场和积分。",
            "一个球员属于一个俱乐部，球员在不同赛季可对应不同的赛季数据，为射手榜和转会扩展预留基础。",
            "一场比赛属于一个主场场馆，关联一个主队俱乐部和一个客队俱乐部，并包含多个票区。",
            "一个票区包含多个座位，票区决定票价；系统根据票区内可售座位的连续情况计算当前最大可接纳连坐数。",
            "一个订单可以包含多张电子票，每张电子票对应一个赛事座位。",
            "一张电子票最多生成一条成功入场记录，但可保留多条失败检票记录用于追溯。",
            "退票申请与订单关联，审核通过后需要同步更新订单、电子票和座位库存状态。",
        ],
    )

    add_heading(doc, "2.3.2 物理结构设计", 3)
    add_table(
        doc,
        ["表名", "主要字段", "说明"],
        [
            ["sys_user", "user_id, username, phone, password, role_id, club_id, status, create_time", "用户与后台账号表，club_id 用于绑定俱乐部或检票员所属俱乐部。"],
            ["sys_role", "role_id, role_name, permission_code, remark", "角色权限表。"],
            ["season_info", "season_id, season_name, start_date, end_date, status", "赛季信息表。"],
            ["round_info", "round_id, season_id, round_name, start_date, end_date", "联赛轮次表。"],
            ["club_info", "club_id, club_name, logo_url, home_city, home_address, home_stadium_id, description", "足球俱乐部信息表，记录俱乐部名称、队标、主场城市、主场地址、主场和简介。"],
            ["player_info", "player_id, club_id, player_name, shirt_no, position, status", "球员信息表，记录俱乐部球员名单，作为射手榜和转会扩展的基础表。"],
            ["player_season_stat", "stat_id, season_id, player_id, club_id, appearances, goals, assists", "球员赛季数据表，记录出场、进球和助攻等扩展统计。"],
            ["coach_info", "coach_id, club_id, coach_name, title, description", "教练信息表，记录俱乐部教练及其职务。"],
            ["club_season_record", "record_id, season_id, club_id, wins, draws, losses, goals_for, goals_against, points", "俱乐部赛季战绩表，记录胜场、平场、负场、进失球和积分。"],
            ["stadium_info", "stadium_id, stadium_name, city, address, capacity, row_count, layout_desc", "场馆信息表，记录主场城市、地址、容量、排数和座位布局说明。"],
            ["stadium_seat_layout", "layout_id, stadium_id, zone_id, row_no, seat_count, start_no, end_no", "主场座位布局表，记录不同票区每一排的座位数量和编号范围。"],
            ["match_info", "match_id, season_id, round_id, home_club_id, away_club_id, stadium_id, match_time, home_score, away_score, status", "比赛信息表，管理员可在比赛结束后导入或维护比分。"],
            ["ticket_zone", "zone_id, match_id, zone_name, price, total_count, remain_count, max_continuous_count", "票区、票价和最大连坐能力表。"],
            ["seat_info", "seat_id, stadium_id, zone_id, row_no, seat_no, center_score, finalized_flag, seat_status", "座位信息表，其中 center_score 用于辅助判断座位是否靠近看台中线，finalized_flag 表示座位所在行是否已最终确定。"],
            ["ticket_order", "order_id, user_id, total_amount, order_status, create_time, expire_time", "订单主表。"],
            ["order_item", "item_id, order_id, match_id, zone_id, seat_id, price", "订单明细表。"],
            ["payment_record", "payment_id, order_id, pay_amount, pay_method, pay_status, pay_time", "支付记录表。"],
            ["e_ticket", "ticket_id, ticket_code, order_id, item_id, ticket_status, enter_time", "电子票表。"],
            ["refund_apply", "refund_id, order_id, reason, audit_status, audit_time, auditor_id", "退票申请表。"],
            ["checkin_record", "record_id, ticket_code, checker_id, check_result, check_time, remark", "检票记录表。"],
        ],
        [1700, 4960, 2700],
    )

    add_heading(doc, "2.4 关键业务规则设计", 2)
    add_table(
        doc,
        ["规则", "设计说明"],
        [
            ["库存防超卖", "订单创建时必须校验座位状态和票区余票，待支付订单锁定系统规划出的座位；支付超时后释放库存。"],
            ["连坐分配", "用户只选择票区和张数，系统优先分配同一排内连续座位；若当前票区无法满足完整连坐，应返回该票区当前最大可接纳连坐数，并提醒用户分开下单。"],
            ["行级座位公布", "连坐必须位于同一排；当某一排尚有优化空间时，用户端不显示具体座位号，只显示票区级信息和最大连坐能力；当该排规划最终确定后，再展示具体座位。"],
            ["场馆布局驱动", "不同俱乐部主场的排数和每排座位数可能不同，连坐算法必须读取主场座位布局数据，不能按固定行宽处理。"],
            ["座位优先级", "在满足连坐的前提下，系统按照下单时间、排号靠前程度和座位居中程度综合排序，越早购买的用户尽量获得更靠前、更居中的连续座位。"],
            ["购票数量限制", "单笔订单最多购买 4 张门票；超过上限时系统拒绝下单并提示用户修改张数。"],
            ["订单状态流转", "订单状态包括待支付、已支付、已取消、已退票；不同状态限制不同操作。"],
            ["电子票唯一性", "每张电子票生成唯一票码，票码与订单明细绑定，检票成功后状态改为已使用。"],
            ["退票限制", "比赛开始后不允许退票；审核通过后释放座位或票区库存，并更新电子票状态。"],
            ["权限控制", "用户只能操作自己的订单；俱乐部只能维护本俱乐部资料、主场布局和主场比赛相关数据；检票员只能使用检票功能；管理员负责全联赛基础数据、赛程票务、比赛成绩、账号权限和系统参数管理。"],
        ],
        [1800, 7560],
    )

    add_heading(doc, "2.5 界面设计说明", 2)
    add_para(
        doc,
        "用户端界面主要包括首页赛事推荐、赛事列表、赛事详情、票区购票、订单确认、支付结果、我的订单和电子票页面。"
        "后台管理界面主要包括俱乐部资料维护、球员教练维护、赛事管理、比赛成绩导入、主场座位布局管理、场馆管理、票区票价管理、订单管理、退票处理、检票记录和销售统计页面。"
        "界面设计应突出购票流程的连续性，使用户能够从赛事浏览自然进入票区选择、张数填写、系统分配座位、下单和查看电子票。"
        "在座位行尚未最终确定时，票区购票页不展示具体座位图，只展示价格、余票、最大可接纳连坐数和购票张数输入框；"
        "当用户所需张数超过当前最大连坐数时，页面提示用户分开下单。"
    )

    add_heading(doc, "2.6 小结", 2)
    add_para(
        doc,
        "足球联赛购票系统围绕职业联赛主客场比赛票务业务建立了较完整的用户端和后台管理功能。系统既包含课程设计要求中的 MVC 分层架构、数据库设计、"
        "功能模块划分和业务流程描述，也具有订单、库存、连坐分配、支付、电子票和检票等较明确的核心业务规则，适合作为软件课程设计项目继续实现。"
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
